#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Literal

from docx import Document


MarkerClass = Literal["terminal", "narrative"]

MARKER_RE = re.compile(r"\[(?:\d{1,3})(?:\s*(?:[,，]|[-–—])\s*\d{1,3})*\]")
LATIN_AUTHOR_RE = r"[A-Z][^\W\d_.'-]*"
ENGLISH_AUTHOR_LED_RE = re.compile(rf"^\s*{LATIN_AUTHOR_RE}(?:\s+et\s+al\.?|等)\s*(\[[^\]]+\])")
AUTHOR_LED_CITATION_RE = re.compile(
    r"(?P<boundary>^|[。！？；;：:，,]\s*)"
    rf"(?P<author>{LATIN_AUTHOR_RE}(?:\s+et\s+al\.?|等))\s*"
    r"(?P<marker>\[(?:\d{1,3})(?:\s*(?:[,，]|[-–—])\s*\d{1,3})*\])"
)
REFERENCE_ENTRY_RE = re.compile(r"^\s*\[\d{1,3}\]\s+\S")


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def iter_document_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        yield from iter_table_paragraphs(table)


def classify_marker(text: str, start: int, end: int) -> MarkerClass:
    return "terminal"


def scan_paragraph_issues(text: str) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    match = ENGLISH_AUTHOR_LED_RE.match(text)
    if match:
        issues.append(
            {
                "kind": "english_author_led_citation",
                "marker": match.group(1),
                "suggestion": f"Rewrite as 文献{match.group(1)}... or 作者在文献{match.group(1)}中...",
                "text": text,
            }
        )
    return issues


def author_led_citation_segments(text: str) -> tuple[str, list[dict[str, str]], list[tuple[str, int]]]:
    rewrites: list[dict[str, str]] = []
    segments: list[tuple[str, int]] = []
    cursor = 0

    for match in AUTHOR_LED_CITATION_RE.finditer(text):
        author_start = match.start("author")
        marker_end = match.end("marker")
        if author_start > cursor:
            segments.append((text[cursor:author_start], cursor))
        marker = match.group("marker")
        author = match.group("author")
        replacement = f"文献{marker}"
        segments.append((replacement, author_start))
        rewrites.append(
            {
                "kind": "english_author_led_citation_rewrite",
                "author": author,
                "marker": marker,
                "replacement": replacement,
            }
        )
        cursor = marker_end

    if cursor < len(text):
        segments.append((text[cursor:], cursor))
    if not rewrites:
        segments = [(text, 0)]
    return "".join(value for value, _source_start in segments), rewrites, segments


def rewrite_author_led_citations(text: str) -> tuple[str, list[dict[str, str]]]:
    rewritten, rewrites, _segments = author_led_citation_segments(text)
    return rewritten, rewrites


def copy_run_format(source, target) -> None:
    target.style = source.style
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.font.name = source.font.name
    target.font.size = source.font.size
    target.font.all_caps = source.font.all_caps
    target.font.small_caps = source.font.small_caps
    target.font.strike = source.font.strike
    target.font.double_strike = source.font.double_strike
    target.font.subscript = source.font.subscript
    target.font.superscript = source.font.superscript
    if source.font.color.rgb is not None:
        target.font.color.rgb = source.font.color.rgb


def run_for_position(run_ranges, position: int):
    for start, end, run in run_ranges:
        if start <= position < end:
            return run
    return run_ranges[-1][2]


def run_ranges_for_paragraph(paragraph):
    run_ranges = []
    cursor = 0
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue
        run_ranges.append((cursor, cursor + len(run_text), run))
        cursor += len(run_text)
    return run_ranges


def replace_paragraph_with_marker_runs(paragraph, text: str) -> int:
    if REFERENCE_ENTRY_RE.match(text):
        return 0

    marker_ranges: list[tuple[int, int, bool]] = []
    styled_markers = 0
    for match in MARKER_RE.finditer(text):
        is_superscript = classify_marker(text, match.start(), match.end()) == "terminal"
        if is_superscript:
            styled_markers += 1
        marker_ranges.append((match.start(), match.end(), is_superscript))
    if not styled_markers:
        return 0

    run_ranges = run_ranges_for_paragraph(paragraph)
    if not run_ranges:
        return 0

    boundaries = {0, len(text)}
    for start, end, _run in run_ranges:
        boundaries.add(start)
        boundaries.add(end)
    for start, end, _is_superscript in marker_ranges:
        boundaries.add(start)
        boundaries.add(end)

    parts: list[tuple[str, bool, object]] = []
    sorted_boundaries = sorted(boundaries)
    for start, end in zip(sorted_boundaries, sorted_boundaries[1:]):
        if start == end:
            continue
        value = text[start:end]
        source_run = run_for_position(run_ranges, start)
        is_superscript = any(
            marker_start <= start and end <= marker_end and marker_superscript
            for marker_start, marker_end, marker_superscript in marker_ranges
        )
        parts.append((value, is_superscript, source_run))

    paragraph.clear()
    for value, superscript, source_run in parts:
        if not value:
            continue
        run = paragraph.add_run(value)
        copy_run_format(source_run, run)
        if superscript:
            run.font.superscript = True
    return styled_markers


def replace_paragraph_text_with_segments(paragraph, segments: list[tuple[str, int]]) -> None:
    run_ranges = run_ranges_for_paragraph(paragraph)
    paragraph.clear()
    for value, source_start in segments:
        if not value:
            continue
        run = paragraph.add_run(value)
        if run_ranges:
            copy_run_format(run_for_position(run_ranges, source_start), run)


def normalize_docx(input_path: str | Path, output_path: str | Path) -> dict[str, object]:
    input_path = Path(input_path)
    output_path = Path(output_path)
    doc = Document(input_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    report: dict[str, object] = {
        "input": str(input_path),
        "output": str(output_path),
        "styled_markers": 0,
        "rewritten_author_led_citations": 0,
        "rewrites": [],
        "issues": [],
    }
    issues: list[dict[str, object]] = []
    rewrites: list[dict[str, object]] = []

    for index, paragraph in enumerate(iter_document_paragraphs(doc)):
        text = paragraph.text
        if not text or "[" not in text:
            continue
        text, paragraph_rewrites, segments = author_led_citation_segments(text)
        if paragraph_rewrites:
            replace_paragraph_text_with_segments(paragraph, segments)
            for rewrite in paragraph_rewrites:
                rewrites.append({"paragraph_index": index, **rewrite})
        for issue in scan_paragraph_issues(text):
            issues.append({"paragraph_index": index, **issue})
        report["styled_markers"] = int(report["styled_markers"]) + replace_paragraph_with_marker_runs(paragraph, text)

    report["issues"] = issues
    report["rewrites"] = rewrites
    report["rewritten_author_led_citations"] = len(rewrites)
    doc.save(output_path)
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Set terminal numeric citation markers as DOCX superscripts.")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--report", type=Path)
    args = parser.parse_args()

    report = normalize_docx(args.input, args.output)
    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    else:
        print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
